# -*- coding: utf-8 -*-
"""
Konwersja docs/production/PRODUCTION_HANDOVER.md -> .docx + .pdf
Bez zmiany treści merytorycznej — tylko formatowanie pod dokument do przekazania.
DOCX: markdown -> HTML (markdown) -> htmldocx.  PDF: własny renderer fpdf2 (Arial/Unicode).
"""
import re, os, markdown
from docx import Document
from docx.shared import Pt, RGBColor
from htmldocx import HtmlToDocx
from fpdf import FPDF
from fpdf.enums import XPos, YPos

SRC  = "docs/production/PRODUCTION_HANDOVER.md"
DOCX = "docs/production/PRODUCTION_HANDOVER.docx"
PDF  = "docs/production/PRODUCTION_HANDOVER.pdf"
ARIAL  = r"C:\Windows\Fonts\arial.ttf"
ARIALB = r"C:\Windows\Fonts\arialbd.ttf"

raw = open(SRC, encoding="utf-8").read()
# Checklisty [ ]/[x] -> ( )/(x) (font-safe; treść zachowana).
md = re.sub(r"\[\s\]", "( )", raw)
md = re.sub(r"\[[xX]\]", "(x)", md)

# Emoji -> etykiety (Arial nie ma glifów emoji). Treść zachowana.
EMOJI = {"🛑":"STOP:", "⚠️":"UWAGA:", "⚠":"UWAGA:", "✅":"[OK]", "🔴":"RYZYKO:",
         "🎉":"", "📍":"", "▸":">", "▾":"v", "🔧":"", "ℹ️":"i:", "ℹ":"i:", "🟢":"[OK]", "🟡":"[~]"}
def clean(s):
    for k, v in EMOJI.items():
        s = s.replace(k, v)
    return s

# ───────────────────────── DOCX (htmldocx) ──────────────────────────────────
body_html = markdown.markdown(md, extensions=["tables", "fenced_code", "sane_lists"])
doc = Document()
n = doc.styles["Normal"]; n.font.name = "Calibri"; n.font.size = Pt(10)
p = doc.add_paragraph(); r = p.add_run("Fresh Market B2B PreConnect"); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(0x0d,0x94,0x88)
p = doc.add_paragraph(); r = p.add_run("Pakiet przekazania do produkcji (Handover) — b2b.freshmarket.eu"); r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x47,0x55,0x69)
doc.add_paragraph()
HtmlToDocx().add_html_to_document(body_html, doc)
doc.save(DOCX)
print("DOCX ->", DOCX, os.path.getsize(DOCX), "B")

# ───────────────────────── PDF (fpdf2) ──────────────────────────────────────
pdf = FPDF(format="A4"); pdf.set_auto_page_break(True, margin=16)
pdf.add_font("B", "", ARIAL); pdf.add_font("B", "B", ARIALB)
pdf.set_margins(16, 14, 16); pdf.add_page()

def H(txt, size, rgb, gap=2.0):
    pdf.ln(gap); pdf.set_font("B", "B", size); pdf.set_text_color(*rgb)
    pdf.multi_cell(0, size*0.5, clean(txt), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
def para(txt, size=9.5, rgb=(30,41,59), bullet=""):
    pdf.set_font("B", "", size); pdf.set_text_color(*rgb)
    pdf.multi_cell(0, size*0.52, ("   "+bullet if bullet else "")+clean(txt), markdown=True, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

lines = md.split("\n"); i = 0
while i < len(lines):
    s = lines[i].strip()
    # blok kodu
    if s.startswith("```"):
        i += 1; code = []
        while i < len(lines) and not lines[i].strip().startswith("```"):
            code.append(lines[i]); i += 1
        i += 1
        pdf.set_font("B", "", 8); pdf.set_text_color(30,41,59); pdf.set_fill_color(244,246,249)
        for c in code:
            pdf.multi_cell(0, 4.0, clean(c) if c.strip() else " ", fill=True, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(1); continue
    # tabela
    if s.startswith("|") and s.endswith("|"):
        tbl = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            tbl.append(lines[i].strip()); i += 1
        rows = [[c.strip() for c in r.strip("|").split("|")] for r in tbl]
        rows = [r for r in rows if not all(c and set(c) <= set("-: ") for c in r)]  # usuń wiersz ---
        if rows:
            ncol = max(len(r) for r in rows)
            pdf.set_font("B", "", 8); pdf.set_text_color(30,41,59); pdf.set_draw_color(148,163,184)
            with pdf.table(markdown=True, text_align="LEFT", first_row_as_headings=True,
                           line_height=4.4, padding=1.3) as table:
                for r in rows:
                    row = table.row()
                    for ci in range(ncol):
                        row.cell(clean(r[ci]) if ci < len(r) else "")
            pdf.ln(1); continue
    # nagłówki
    if s.startswith("# "):   H(s[2:], 16, (13,148,136), 3); pdf.ln(0.5); i += 1; continue
    if s.startswith("## "):  H(s[3:], 12.5, (15,23,42), 3); i += 1; continue
    if s.startswith("### "): H(s[4:], 10.5, (51,65,85), 1.5); i += 1; continue
    # hr
    if s in ("---","***","___"):
        pdf.set_draw_color(203,213,225); y = pdf.get_y()+1
        pdf.line(pdf.l_margin, y, pdf.w-pdf.r_margin, y); pdf.ln(3); i += 1; continue
    # cytat
    if s.startswith(">"):
        q = []
        while i < len(lines) and lines[i].strip().startswith(">"):
            q.append(lines[i].strip().lstrip(">").strip()); i += 1
        pdf.set_font("B", "", 9); pdf.set_text_color(71,85,105); pdf.set_fill_color(248,250,252)
        pdf.multi_cell(0, 4.6, clean(" ".join(q)), markdown=True, fill=True, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(1); continue
    # listy
    if s.startswith("- ") or s.startswith("* "): para(s[2:], 9, bullet="- "); i += 1; continue
    m = re.match(r"^(\d+)\.\s+(.*)", s)
    if m: para(m.group(2), 9, bullet=m.group(1)+". "); i += 1; continue
    if s == "": pdf.ln(2); i += 1; continue
    para(s); i += 1

pdf.output(PDF)
print("PDF  ->", PDF, os.path.getsize(PDF), "B")
